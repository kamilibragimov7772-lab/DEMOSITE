"""Извлекает текст из ТЗ.docx в чистый markdown."""
import sys
from pathlib import Path
from docx import Document

SRC = Path(r"C:\Users\kamil\OneDrive\Документы\TERMY-site-polish\02_github_repo\bot-ponton\ТЗ_на_лучший_B2B_сайт_TERMY_2026 (1).docx")
DST = Path(r"C:\Users\kamil\OneDrive\Документы\TERMY-site-polish\03_content_draft\00_tz_extracted.md")

doc = Document(str(SRC))

lines = ["# ТЗ на лучший B2B-сайт TERMY 2026", "", f"*Источник: {SRC.name}*", ""]

for p in doc.paragraphs:
    text = p.text.strip()
    if not text:
        continue
    style_name = getattr(p.style, "name", None) if p.style else None
    style = (style_name or "").lower()
    if "heading 1" in style:
        lines.append(f"\n## {text}\n")
    elif "heading 2" in style:
        lines.append(f"\n### {text}\n")
    elif "heading 3" in style:
        lines.append(f"\n#### {text}\n")
    elif "heading" in style:
        lines.append(f"\n##### {text}\n")
    elif "list" in style or p.text.lstrip().startswith(("•", "-", "*", "·")):
        lines.append(f"- {text}")
    else:
        lines.append(text)

# Tables
if doc.tables:
    lines.append("\n\n---\n\n## Таблицы\n")
    for i, t in enumerate(doc.tables, 1):
        lines.append(f"\n### Таблица {i}\n")
        for row in t.rows:
            cells = [c.text.strip().replace("\n", " ") for c in row.cells]
            lines.append("| " + " | ".join(cells) + " |")
        lines.append("")

DST.write_text("\n".join(lines), encoding="utf-8")
print(f"Saved {DST}")
print(f"  paragraphs: {len(doc.paragraphs)}")
print(f"  tables: {len(doc.tables)}")
print(f"  lines: {len(lines)}")
print(f"  chars: {sum(len(l) for l in lines)}")
